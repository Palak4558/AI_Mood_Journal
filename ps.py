
from docx import Document

# Create a new Document
doc = Document()

# Add a title
doc.add_heading("Problem Statement", level=1)

# Add some text
doc.add_paragraph("Mental health probelms has limited olution, there are only two major solution to it: one is a totally clinical approach and the other one is negelcting it. There can be a mid-way in it , where we get the best of both the worlds - that is through a AI Mental Health Journal")

# Save the document
doc.save("Problem_statement.docx")
